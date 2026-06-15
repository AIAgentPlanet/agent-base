from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Agent Base智能体互联网交互治理研究报告.docx")

BLUE = "2E5E8C"
DARK = "17324D"
LIGHT = "EAF1F7"
PALE = "F5F7FA"
GRAY = "666666"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, size=10.5, bold=False, color="222222", font="Arial Unicode MS"):
    run.font.name = font
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:eastAsia"), font)
    fonts.set(qn("w:cs"), font)
    fonts.set(qn("w:ascii"), font)
    fonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", bold_lead=None, align=None, after=6, before=0, first_indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.35
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(21)
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        style_run(r1, bold=True, color=DARK)
        r2 = p.add_run(text[len(bold_lead):])
        style_run(r2)
    else:
        style_run(p.add_run(text))
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    for r in p.runs:
        style_run(r)
    if not p.runs:
        style_run(p.add_run(text))
    else:
        p.runs[0].text = text
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    style_run(p.add_run(text))
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        set_cell_shading(hdr.cells[i], BLUE)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        style_run(p.add_run(header), size=9.5, bold=True, color=WHITE)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            if row_idx % 2:
                set_cell_shading(cells[i], PALE)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if i == 0 and len(value) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_run(p.add_run(value), size=9)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    style_run(p.add_run(title + "  "), size=10, bold=True, color=DARK)
    style_run(p.add_run(text), size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_architecture(doc):
    layers = [
        ("用户与智能体入口层", "CLI、IDE 插件、多模态入口、第三方 AI Agent"),
        ("意图与任务编排层", "提示词与 Skill 定义、任务拆解、能力选择、用户确认"),
        ("可信交互协议层", "MCP（规划建设）+ ATH M1—M5（已在 user-service 落地）"),
        ("治理与服务层", "DID 身份、双向握手、会话密钥、Scope、请求完整性、吊销与签名审计"),
        ("数据与能力层", "PostgreSQL 审计链/Outbox、Redis 会话/重放保护、用户与 OAuth API"),
    ]
    for idx, (name, detail) in enumerate(layers):
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        set_table_geometry(table, [8200])
        cell = table.cell(0, 0)
        set_cell_shading(cell, BLUE if idx in (0, 2, 4) else LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        style_run(p.add_run(name + "\n"), size=10.5, bold=True, color=WHITE if idx in (0, 2, 4) else DARK)
        style_run(p.add_run(detail), size=9.5, color=WHITE if idx in (0, 2, 4) else "333333")
        if idx < len(layers) - 1:
            p2 = doc.add_paragraph("↓")
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_after = Pt(0)
            style_run(p2.runs[0], size=12, bold=True, color=BLUE)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1.0)
sec.bottom_margin = Inches(1.0)
sec.left_margin = Inches(1.0)
sec.right_margin = Inches(1.0)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial Unicode MS"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
normal._element.rPr.rFonts.set(qn("w:cs"), "Arial Unicode MS")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.35
for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, DARK, 8, 4),
]:
    st = styles[name]
    st.font.name = "Arial Unicode MS"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    st._element.rPr.rFonts.set(qn("w:cs"), "Arial Unicode MS")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

# Cover
for _ in range(5):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(p.add_run("智能体互联网治理"), size=26, bold=True, color=DARK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
style_run(p.add_run("交互技术架构与产业实践研究报告"), size=22, bold=True, color=BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(42)
style_run(p.add_run("——基于 Agent Base 开源项目的治理实践"), size=14, color=GRAY)
add_callout(doc, "项目定位", "面向 AI 原生 Web 应用开发的底座框架，以提示词与技能、MCP/CLI 接入、标准微服务模板和 ATH 可信握手协议支撑智能体安全调用业务能力。")
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(p.add_run("Agent Base 开源项目团队"), size=12, bold=True, color=DARK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(p.add_run("2026 年 6 月"), size=11, color=GRAY)
doc.add_page_break()

# Front matter
add_heading(doc, "报告摘要", 1)
add_para(doc, "大模型与智能体正从信息问答工具转变为连接用户、应用、数据、工具与设备的行动入口。交互链路由“人—系统”扩展为“人—智能体—工具—服务—数据”的多主体协作网络，身份、授权、上下文、调用结果与责任边界因此成为新的治理对象。")
add_para(doc, "本报告依据《智能体互联网治理交互技术架构与产业实践研究报告（2026）》大纲，以 Agent Base 项目为实践样本，系统分析智能体互联网交互技术、产业生态、风险类型与治理能力，并总结 ATH（Agent Trust Handshake）协议在用户服务中的实现经验。")
add_para(doc, "研究认为，智能体交互治理应形成覆盖接入前、运行中和事后环节的闭环：接入前实现主体可识别、工具可登记、能力可发现；运行中实现请求可校验、权限可约束、会话可绑定、异常可阻断；事后实现链路可还原、证据不可静默篡改、责任可追溯。Agent Base 已完成 ATH M1—M5 阶段改造，形成 DID/身份文档验证、双向挑战签名、P-256 ECDH 与 HKDF 会话密钥、HMAC 请求完整性、签名审计哈希链、密钥轮换、事务型 Outbox 和外部锚定能力。")
add_callout(doc, "核心结论", "Agent Base 已从“OAuth 上增加 Agent 参数”的早期实现，演进为具备强身份、短期会话密钥、逐请求完整性和可验证存证的可信交互网关原型。生产化重点已转向真实基础设施联调、资源/参数级策略、限流与风控、MCP 工具生态和第三方互操作测评。")

add_heading(doc, "项目概况", 1)
add_table(doc, ["项目要素", "内容"], [
    ("项目名称", "Agent Base"),
    ("项目性质", "AI 原生 Web 应用开发底座框架"),
    ("核心模块", "prompts、skills、mcp、cli、services"),
    ("已落地能力", "user-service、JWT/OAuth 2.0、ATH M1—M5、DID 双向认证、会话密钥、请求完整性、签名审计链、外部锚定"),
    ("规划能力", "MCP Server 核心实现、CLI 发布、资源/参数级动态策略、文件/消息/权限服务、多语言 SDK"),
    ("技术栈", "Go、Gin、PostgreSQL、Redis、JWT、bcrypt、OAuth 2.0"),
], [1800, 7560])

add_heading(doc, "目录", 1)
for item in [
    "一、智能体互联网交互发展概述",
    "二、智能体互联网交互风险分析",
    "三、智能体互联网交互治理能力体系建设",
    "四、交互治理产业实践与典型应用场景",
    "五、智能体互联网交互治理发展展望",
    "附录：Agent Base 能力状态与证据索引",
]:
    add_para(doc, item, first_indent=False, after=3)
doc.add_page_break()

# Chapter 1
add_heading(doc, "一、智能体互联网交互发展概述", 1)
add_heading(doc, "（一）交互技术发展", 2)
add_heading(doc, "1. 交互技术架构与关键能力", 3)
add_para(doc, "智能体互联网交互并非单一模型能力，而是一套把意图理解、任务编排、工具发现、身份认证、授权决策、接口调用、数据流转和结果反馈串联起来的系统工程。与传统 API 集成相比，智能体具有更强的自主决策和跨系统组合能力，治理控制点必须从单个接口扩展到完整调用链。")
add_architecture(doc)
add_para(doc, "Agent Base 将业务知识沉淀在 prompts/ 与 skills/，将标准化调用入口规划在 mcp/ 与 cli/，并以 services/ 提供可复用的微服务模板。当前 user-service 已形成“用户身份—OAuth 授权—ATH 智能体授权—代理调用”的核心闭环。")
add_table(doc, ["能力域", "治理目标", "Agent Base 对应能力", "状态"], [
    ("用户入口", "统一身份入口与用户授权", "注册、登录、JWT、OAuth 授权页", "已实现"),
    ("任务编排", "将意图约束为可审核调用", "Skill 定义；ATH proxy 统一入口", "部分实现"),
    ("工具调用", "可发现、可授权、可校验、可阻断", "ATH 发现文档、Provider/Scope、HMAC 请求完整性", "已实现核心"),
    ("协议适配", "兼容通用协议并扩展 Agent 强身份", "OAuth 2.0 + ATH v0.1 网关模式", "已实现核心"),
    ("端云协同", "本地入口与云端服务可信连接", "CLI、MCP Server", "规划中"),
    ("数据流转", "最小化、隔离、可验证追溯", "短期会话密钥、Redis、签名审计链、外部锚定", "已实现基础"),
], [1200, 2200, 4400, 1560])

add_heading(doc, "2. 交互协议与工具调用能力", 3)
add_para(doc, "MCP 面向模型与外部上下文、资源及工具之间的标准连接，可降低智能体接入不同业务能力的适配成本。Agent Base 已为 mcp/ 预留实现位置，目标是将用户、文件、消息和权限等服务封装为平台无关的工具接口。现阶段该模块仍处于建设阶段，不应视为已经形成生产级 MCP 服务。")
add_para(doc, "ATH 是项目在 user-service 中实现的智能体可信交互协议。其设计思路是在 OAuth 2.0 用户授权基础上增加 Agent 身份文档、双向身份挑战、Provider/Scope 申请、会话密钥、逐请求完整性和签名存证，从而把“谁代表谁、双方是否真实、可调用什么、本次请求是否被篡改、如何追溯”纳入协议流程。")
add_table(doc, ["步骤", "接口/动作", "主要治理作用"], [
    ("1 发现", "GET /.well-known/ath.json", "发布端点、算法、能力和审计/锚定地址"),
    ("2 注册", "POST /api/v1/ath/agents/register", "解析身份文档并验证 ES256 Attestation、声明与 JTI"),
    ("3 双向握手", "POST /handshakes；POST /proof", "服务端挑战签名、Agent 身份证明、原子状态迁移"),
    ("4 会话建钥", "P-256 ECDH + HKDF-SHA256", "双方独立派生 32 字节短期会话密钥"),
    ("5 用户授权", "ATH authorize + OAuth consent", "绑定已验证握手并计算 Scope 交集"),
    ("6 Token 交换", "POST /api/v1/ath/token", "将 Token 与 Agent、Provider、Session、Handshake 绑定"),
    ("7 代理调用", "POST /api/v1/ath/proxy", "校验 HMAC、时间窗、nonce、Provider 与 Scope"),
    ("8 审计锚定", "audit query/verify/head/anchor", "验证签名哈希链并可靠投递独立锚点"),
], [1100, 3300, 4960])

add_heading(doc, "3. 技术成熟度与发展难点", 3)
add_para(doc, "从项目实现看，ATH 已完成从身份注册到外部存证的纵向闭环。服务端可从公共 HTTPS Agent Identity Document 或 did:web 文档取得 P-256 公钥，验证 ES256 Attestation 的签名、iss、sub、aud、iat、exp 和 jti；握手阶段通过双方签名和 ECDH/HKDF 建立短期会话密钥；代理请求使用 HMAC 绑定 Token、方法、路径、正文摘要和握手，并以 Redis 原子 nonce 防止重放；关键事件进入签名哈希链和事务型 Outbox。")
add_table(doc, ["难点", "当前表现", "治理影响", "建议"], [
    ("跨平台身份互认", "已支持公共 HTTPS 身份文档与 did:web", "仍需与更多 Agent 平台互操作", "形成测试向量、兼容配置与第三方互认测试"),
    ("密钥生产托管", "支持本地密钥环和通用 HTTPS KMS 网关", "尚非云厂商原生 KMS SDK 集成", "完善 KMS 网关鉴权、可用性、审批与灾备"),
    ("上下文管理", "当前重点在 Token/Session，长期记忆治理未覆盖", "污染与跨域泄露风险仍在", "增加来源标记、隔离域、保留期限和清除机制"),
    ("工具代理粒度", "已校验 Provider/Scope、方法、路径和请求完整性", "资源与业务参数仍可能过度授权", "引入资源级 ABAC、参数策略和高风险二次确认"),
    ("端云协同", "CLI/MCP 尚在规划", "无法验证复杂设备与本地数据边界", "先建设最小可用连接器和本地确认机制"),
    ("生产环境验证", "单元、竞态、编译和静态检查已通过", "真实数据库锁与外部 Webhook 尚需联调", "建设 PostgreSQL/Redis/锚点集成测试环境"),
], [1500, 2500, 2400, 2960])

add_heading(doc, "（二）交互产业与应用发展", 2)
add_heading(doc, "1. 交互产业生态发展情况", 3)
add_para(doc, "智能体交互生态通常由模型与平台方、云服务方、终端厂商、工具服务商、应用服务商、安全机构和行业用户共同组成。平台提供模型与编排，云侧提供算力和托管工具，终端承载用户入口与本地数据，工具和应用方提供可执行能力，安全机构承担检测、审计和认证，行业用户则决定业务授权与责任边界。")
add_para(doc, "Agent Base 的定位更接近“开发底座 + 可信接入网关”：通过标准服务模板保证代码一致性，通过 skills 沉淀领域知识，通过 ATH 建立 Agent 到业务系统的授权通道。其价值不在于替代模型平台，而在于为企业自有系统提供可控的智能体接入层。")

add_heading(doc, "2. 交互应用场景分布", 3)
add_table(doc, ["场景", "交互方式", "主要治理要求"], [
    ("企业办公", "Agent 调用用户、权限、消息和文档系统", "身份代理、最小授权、敏感操作确认"),
    ("开发工具", "IDE/CLI Agent 调用代码、配置和部署能力", "本地数据边界、命令白名单、凭证隔离"),
    ("云服务", "多 Agent 调用云端 API 与微服务", "租户隔离、动态策略、链路追踪"),
    ("安全运营", "Agent 查询资产并执行响应动作", "高风险动作双确认、可回滚、强审计"),
    ("政务与金融", "Agent 辅助办理、审核与查询", "实名授权、数据最小化、人工复核、留痕"),
], [1500, 3900, 3960])

add_heading(doc, "3. 交互产业发展需求", 3)
for text in [
    "统一接入：降低不同模型、插件、API 和企业系统之间的重复适配成本。",
    "可信授权：把用户意图转换为明确、短期、可撤销的机器权限。",
    "可组合治理：在多工具、多步骤任务中持续传播身份、权限和审计上下文。",
    "开放生态：通过标准能力描述和测试规范促进工具互认，同时避免无边界开放。",
    "工程落地：提供模板、SDK、测试脚本和部署方案，使治理要求可被开发团队直接采用。",
]:
    add_bullet(doc, text)

add_heading(doc, "（三）交互监管政策与标准", 2)
add_heading(doc, "1. 政策与合规要求", 3)
add_para(doc, "智能体交互在中国境内落地时，应结合《网络安全法》《数据安全法》《个人信息保护法》以及生成式人工智能服务相关管理要求，落实合法正当必要、目的限定、最小授权、安全保障和个人权利响应。若智能体处理个人信息、跨境数据、重要数据或面向公众提供生成式服务，还需根据具体业务主体、部署方式和数据类型开展专项合规评估。")
add_para(doc, "对 Agent Base 而言，合规要求应转化为工程控制：将 Provider/Scope 与处理目的对应；在授权页面说明访问对象与期限；对日志执行脱敏和保留期限管理；对 Token、密钥和 Session 实施安全存储；为删除、更正、撤回授权和事件响应提供接口。")
add_heading(doc, "2. 标准与测评需求", 3)
add_para(doc, "当前急需围绕交互术语、参考架构、身份声明、能力发现、授权范围、敏感操作确认、审计字段、错误码和测试方法形成可互操作规范。标准化不宜只描述抽象原则，还应给出最小数据结构、协议状态机、负向测试和证据输出要求。")

# Chapter 2
add_heading(doc, "二、智能体互联网交互风险分析", 1)
add_heading(doc, "（一）主体身份与身份互认风险", 2)
add_para(doc, "交互链路中的主体至少包括用户、Agent、开发者、工具提供方、业务服务和设备。若 Agent 身份仅依赖可伪造的名称或客户端参数，服务端无法判断调用者是否为登记主体；若不同平台对用户和 Agent 的绑定方式不一致，权限在跨平台传递时也可能失真。")
add_para(doc, "Agent Base 已将“声明身份”“验证身份”和“会话身份”分离：注册阶段验证身份文档、ES256 签名和标准声明；握手阶段由服务端与 Agent 分别对挑战上下文签名；授权和 Token 阶段强制绑定已验证 handshake_id。密钥环可发布历史验证方法并切换 active key，远程 KMS 网关可避免业务进程直接持有生产私钥。")

add_heading(doc, "（二）权限越界与工具调用风险", 2)
add_para(doc, "智能体可能因提示理解偏差、工具描述不准确、Scope 过宽或代理规则粗放而执行超出用户意图的操作。写入、删除、转账、发布、创建凭证等动作尤其需要把“可调用接口”进一步细化为“可操作资源、可使用参数、可执行次数和有效时间”。")
add_para(doc, "项目已实现 Agent 获批 Scope、请求 Scope 与用户同意 Scope 的交集，并在 ath/proxy 校验 Token、Provider、Scope、方法、路径和 HMAC 请求签名。该机制能发现正文、路径和上下文绑定被修改，也能拒绝过期时间戳和重复 nonce；生产化仍需引入资源级策略、参数策略、速率限制、敏感动作二次确认和幂等/回滚机制。")

add_heading(doc, "（三）数据流转与上下文污染风险", 2)
add_para(doc, "智能体使用的会话、长期记忆、知识库、工具返回和端侧文件均可能携带敏感数据或恶意指令。工具返回内容若未经信任分级直接进入下一轮推理，可能形成间接提示注入；不同用户或任务的记忆混用，则可能造成跨域泄露。")
add_para(doc, "Agent Base 当前通过 Redis 隔离 ATH Session 与握手状态，Token 有效期受安全会话剩余时间约束，请求 nonce 使用原子一次性写入；会话密钥仅用于本次握手关联的完整性校验。后续仍应增加数据分类、上下文来源标签、工具结果净化、跨租户隔离、记忆保留与删除策略，以及对日志中 Token、code、个人信息的脱敏。")

add_heading(doc, "（四）交互链路与责任边界风险", 2)
add_para(doc, "多步骤任务可能跨越模型平台、Agent、代理网关和多个业务工具。若仅记录最终 API 请求，难以还原用户原始意图、Agent 决策、授权范围和工具结果之间的关系。链路异常时，平台方、Agent 开发者、工具方和业务运营者也可能相互推诿。")
add_para(doc, "项目已形成结构化审计事件模型，以 event_id、全局 sequence、previous_hash、payload_hash、record_hash、signing_key_id 和 ES256 signature 构成连续证据链。PostgreSQL 触发器拒绝审计记录更新和删除；审计记录与 Outbox 在同一事务写入，外部锚定 worker 使用幂等键、SKIP LOCKED、多实例领取、指数退避和崩溃锁恢复进行可靠投递。")

add_table(doc, ["风险", "发生可能性", "影响", "现有控制", "剩余风险"], [
    ("Agent 身份伪造", "低中", "高", "身份文档、ES256 Attestation、双向挑战、JTI", "身份文档托管与密钥运营仍是信任点"),
    ("Scope 越权", "中", "高", "三方 Scope 取交集、Token/Provider/路径校验", "资源/参数级策略与限流不足"),
    ("请求篡改或重放", "低中", "高", "会话 HMAC、时间窗、Redis 原子 nonce", "客户端密钥保护和时钟治理需加强"),
    ("提示注入影响工具调用", "中高", "高", "统一代理、方法/路径和完整性校验", "上下文信任分级与内容检测不足"),
    ("审计证据被篡改", "低中", "高", "签名哈希链、追加触发器、外部锚定 Outbox", "数据库整体回滚及锚点可用性仍需治理"),
    ("规划能力被误认为已交付", "中", "中", "README 路线图说明", "需建立能力状态清单和发布门禁"),
], [1760, 1200, 1000, 3100, 2300])

# Chapter 3
add_heading(doc, "三、智能体互联网交互治理能力体系建设", 1)
add_heading(doc, "（一）交互治理总体能力框架", 2)
add_para(doc, "建议以“可识别、可授权、可管控、可审计、可追责、可协同”为总体目标，覆盖主体、对象、链路和结果四类治理对象，并按接入前、运行中、事后三个阶段部署控制。该框架既适用于 Agent Base，也可作为企业建设智能体交互网关的通用参考。")
add_table(doc, ["阶段", "核心问题", "关键能力", "主要证据"], [
    ("接入前", "谁可接入、可调用什么", "身份审核、工具准入、能力目录、权限分级、协议适配", "登记记录、密钥证明、Provider/Scope 清单"),
    ("运行中", "本次调用是否符合用户意图和策略", "Token 校验、策略决策、参数检查、限流、确认、隔离、阻断", "请求 ID、策略结果、确认记录、异常码"),
    ("事后", "发生了什么、由谁负责、如何改进", "日志审计、链路还原、吊销处置、复盘与规则优化", "审计事件、撤销记录、事件报告、测试结果"),
], [1200, 2800, 3200, 2160])

add_heading(doc, "（二）接入前治理能力", 2)
for text in [
    "主体身份审核：登记 Agent 开发者、Agent ID、身份文档、回调地址和用途；强制验证 ES256 签名、签发者、受众、时间与唯一标识。",
    "工具准入：对 Provider 开展安全评估，登记接口、数据类型、风险等级、依赖关系与责任人。",
    "能力目录：通过发现文档发布端点、Scope、握手算法、审计和锚定接口，并以 DID 文档发布当前及历史验证密钥。",
    "权限分级：区分读取、写入、管理和高风险操作；默认拒绝未声明能力。",
    "协议适配：对 OAuth、MCP、ATH 等协议建立统一身份和审计映射，避免适配层丢失安全语义。",
]:
    add_bullet(doc, text)
add_callout(doc, "已完成改造", "M1—M5 已依次完成强制 Attestation 验签、双向挑战、ECDH/HKDF 会话密钥、逐请求 HMAC、签名审计哈希链、密钥轮换、远程 KMS 网关适配、事务型 Outbox 与外部锚定。")

add_heading(doc, "（三）运行中管控能力", 2)
add_table(doc, ["控制点", "当前基础", "增强要求"], [
    ("请求校验", "ATH Token、HMAC、时间窗、nonce、正文摘要", "增加请求体大小、字段白名单和内容类型策略"),
    ("权限控制", "Provider/Scope/方法/路径与握手绑定", "扩展到资源、参数和环境条件"),
    ("敏感确认", "OAuth 用户授权", "对单次高风险动作展示影响范围并二次确认"),
    ("调用限流", "Redis 基础设施可用", "按 Agent、用户、Scope、接口与风险等级限流"),
    ("上下文隔离", "握手、ATH Session 和 nonce 独立存储并过期", "增加租户、任务、记忆和工具结果隔离"),
    ("异常阻断", "错误码与 Token 撤销", "接入风险评分、熔断、降级和人工接管"),
], [1700, 3200, 4460])

add_heading(doc, "（四）事后审计与责任追溯能力", 2)
add_para(doc, "项目已将普通运行日志升级为可验证治理证据。审计服务对规范化事件载荷计算 SHA-256 摘要，将每条记录与前序哈希连接，并由当前网关身份密钥签名。验证接口能够检查序号连续性、前序链接、载荷摘要、记录哈希与签名；公开链头可供独立监控方留存。")
add_para(doc, "为解决“数据库写入成功但外部投递失败”的双写问题，项目采用 Transactional Outbox：审计记录和待投递锚点在同一事务提交。锚定 worker 以 event_id 作为 Idempotency-Key，携带序号和记录哈希，支持多实例安全领取、失败退避、状态查询和人工重试。")

add_heading(doc, "（五）测试评估与标准化能力", 2)
add_para(doc, "项目已提供 scripts/test-ath.sh，并建立身份文档、Attestation、握手、状态机、密钥环、远程签名、请求完整性、重放防护、审计链和锚定 worker 的自动化测试。当前针对性单元测试、竞态检查、编译检查、构建和静态检查已通过；真实 PostgreSQL、Redis 与外部 Webhook 的端到端联调仍是生产验收项。")
add_table(doc, ["测试域", "关键用例", "通过标准"], [
    ("身份与握手", "伪造签名、错误声明、弱 nonce、陈旧证明、挑战篡改、状态冲突", "非法身份拒绝，状态仅可原子迁移"),
    ("会话与授权", "错误共享秘密、超范围 Scope、过期握手、授权码复用", "密钥一致、权限不扩大、状态机不可绕过"),
    ("代理调用", "正文/路径篡改、陈旧时间戳、nonce 重放、未登记路径", "默认拒绝并返回稳定错误码"),
    ("数据保护", "日志泄露 Token、跨用户 Session、工具结果注入", "敏感信息不明文泄露，上下文严格隔离"),
    ("审计与锚定", "载荷/链路篡改、错误 KMS 签名、Webhook 失败、重复投递", "验证失败可定位，投递可恢复且幂等"),
], [1600, 4800, 2960])

# Chapter 4
add_heading(doc, "四、交互治理产业实践与典型应用场景", 1)
add_heading(doc, "（一）Agent Base 工具调用与协议治理实践", 2)
add_para(doc, "在企业内部场景中，user-service 可作为智能体访问用户与 OAuth 能力的认证授权入口。Agent 先以身份文档和 Attestation 建立长期身份，再与网关完成双向挑战并派生短期会话密钥；用户确认授权范围后，Agent 使用绑定 handshake_id 的 ATH Token，经带 HMAC 的 ath/proxy 请求调用受控接口。")
add_para(doc, "实践价值主要体现在四方面：第一，将 Agent 身份与普通用户 Token 区分，并证明通信双方身份；第二，以 Scope 和统一代理减少内部接口直接暴露；第三，以短期会话密钥和一次性 nonce 防止请求篡改与重放；第四，以签名哈希链和外部锚定形成独立于普通应用日志的治理证据。")
add_table(doc, ["里程碑", "完成能力", "代表性验收"], [
    ("M1 强身份", "身份文档、公钥绑定、ES256 Attestation 完整验签", "错误签名/声明/算法降级均拒绝"),
    ("M2 双向握手", "服务端挑战签名、Agent proof、Redis 原子状态机", "篡改挑战、弱 nonce、陈旧 proof 均拒绝"),
    ("M3 安全会话", "P-256 ECDH、HKDF-SHA256、Token/handshake 绑定", "双方派生一致，过期会话无法换取或使用 Token"),
    ("M4 完整性存证", "HMAC 请求签名、重放防护、ES256 审计哈希链", "正文/路径/nonce 篡改可检测，审计链可验证"),
    ("M5 生产增强", "密钥轮换、远程 KMS、Outbox、外部锚定", "历史公钥可验证，投递失败可退避恢复"),
], [1450, 4700, 3210])

add_heading(doc, "（二）用户入口与多模态交互治理实践", 2)
add_para(doc, "当前项目的成熟入口以 API、Swagger 和技能说明为主，多模态入口尚未落地。后续接入 IDE、CLI、语音或图像入口时，应保持统一身份和授权语义：入口只负责采集意图，不得自行扩大权限；对识别置信度低或涉及敏感资源的请求，应回到可读、可确认的交互界面。")

add_heading(doc, "（三）端云协同与设备交互治理实践", 2)
add_para(doc, "规划中的 CLI 与 MCP Server 可作为本地开发环境和云端服务之间的连接器。建议采用“本地最小采集、云端最小授权、敏感动作本地确认”的原则：代码和配置仅在任务必要范围内上传；云端 Token 不写入普通日志；执行部署、删除、凭证创建等命令前展示目标环境与影响范围。")

add_heading(doc, "（四）行业场景交互治理实践", 2)
add_table(doc, ["行业/场景", "Agent Base 适配方式", "关键控制"], [
    ("企业办公", "以用户服务和后续权限/消息服务支撑成员管理与审批", "组织身份绑定、审批链、写操作确认"),
    ("安全运营", "Agent 经 ATH 访问资产和用户信息并触发处置", "只读/处置 Scope 分离、双人复核、回滚"),
    ("云原生平台", "为多微服务提供统一 Agent 接入网关", "租户隔离、服务目录、限流与链路追踪"),
    ("开发工具", "skills + MCP/CLI 对接 IDE Agent", "仓库边界、命令策略、密钥隔离"),
    ("政务金融", "作为受控接口层连接既有业务系统", "实名授权、最小数据、人工复核、长期审计"),
], [1500, 4300, 3560])

add_heading(doc, "（五）交互安全评测与运营实践", 2)
for text in [
    "建立能力台账：区分已实现、实验性、规划中和已废弃能力，避免文档与交付状态不一致。",
    "持续负向测试：在 CI 中覆盖伪造身份、挑战篡改、重放、越权 Scope、链路篡改、错误 KMS 签名和锚点故障。",
    "运营监测：按 Agent、用户、Scope、Provider 和签名密钥统计调用量、拒绝率、审计链状态及 Outbox 积压。",
    "事件响应：支持 Token 吊销、Agent 冻结、密钥切换、锚定补投、证据导出和影响范围定位。",
    "版本治理：发现文档、DID 验证方法、Scope 语义和代理规则变更需版本化并提供兼容期。",
]:
    add_bullet(doc, text)

# Chapter 5
add_heading(doc, "五、智能体互联网交互治理发展展望", 1)
add_heading(doc, "（一）交互治理技术发展展望", 2)
add_table(doc, ["方向", "发展判断", "Agent Base 路线"], [
    ("协议化", "身份、能力发现、授权和审计将形成机器可读协议", "固化 ATH 测试向量并兼容 OAuth/MCP 身份语义"),
    ("接口化", "治理能力将以网关、SDK 和策略 API 提供", "建设 MCP Server 与多语言 SDK"),
    ("策略化", "静态 Scope 将升级为上下文感知动态授权", "引入 RBAC/ABAC/风险策略引擎"),
    ("可信化", "远程证明、密钥硬件保护与可验证日志逐步应用", "对接云原生 KMS/HSM、可信时间戳与透明日志"),
    ("自动化", "安全测试、注册、发布和审计证据自动生成", "把 ATH 集成测试和锚点恢复演练纳入 CI/CD"),
    ("轻量化", "中小团队需要低门槛默认安全模板", "提供一键部署、示例策略和参考服务"),
], [1200, 3900, 4260])

add_heading(doc, "（二）交互治理产业发展展望", 2)
add_para(doc, "交互网关将成为连接模型平台与企业系统的关键基础设施，其产品边界会从认证代理扩展到工具目录、策略决策、上下文安全、调用监测和事件响应。围绕工具安全评级、Agent 身份服务、交互安全评测、审计托管和行业策略包将形成新的专业服务。")
add_para(doc, "Agent Base 可沿“开源开发底座—可信交互网关—行业能力模板”路径演进：以 user-service 验证身份和授权，以 MCP/CLI 扩大工具连接，以 skills 沉淀行业规范，再通过测试与部署模板降低企业采用成本。")

add_heading(doc, "（三）监管与标准化展望", 2)
for text in [
    "参考架构：明确用户、Agent、编排器、工具、网关、数据源和设备的角色及信任边界。",
    "身份与授权：规定 Agent 身份声明、密钥证明、用户委托、Scope 语义、有效期和吊销机制。",
    "接口协议：统一能力发现、错误码、敏感操作确认、审计上下文传播和版本协商。",
    "测试方法：建立覆盖身份、授权、工具调用、上下文污染、数据保护和审计追溯的用例库。",
    "实践指南：针对企业办公、开发工具、云服务和高风险行业形成分级落地要求。",
    "评估认证：依据可验证证据评估交互治理成熟度，避免只审查制度文件而忽略代码和运行效果。",
]:
    add_number(doc, text)

add_heading(doc, "（四）项目实施路线建议", 2)
add_table(doc, ["优先级", "建设周期", "重点任务", "验收证据"], [
    ("P0", "近期", "真实 PostgreSQL/Redis/Webhook 联调；Outbox 故障恢复演练；修复存量 SQLMock 测试", "集成测试报告、恢复时间、全量测试通过"),
    ("P1", "中期", "资源/参数级 ABAC；Agent/Scope 限流；敏感动作二次确认；运营指标", "策略版本、压测、越权测试和仪表盘"),
    ("P2", "中长期", "云 KMS/HSM 原生适配；MCP Server、CLI、多服务模板与 SDK", "密钥演练、可运行示例、跨平台互操作测试"),
    ("P3", "持续", "行业策略包、第三方测评、可信时间戳/透明日志、标准协同", "行业试点、测评报告、锚定证明和标准提案"),
], [1000, 1300, 4500, 2560])

# Appendix
add_heading(doc, "附录：Agent Base 能力状态与证据索引", 1)
add_table(doc, ["能力", "代码/材料证据", "状态判断"], [
    ("项目定位与模块", "README.MD；prompts/、skills/、mcp/、cli/、services/", "目录与路线图明确"),
    ("用户与 OAuth", "services/user-service/internal/handler；routers/oauth.go；pkg/oauth", "已实现"),
    ("ATH 发现与路由", "routers/ath.go；pkg/ath/discovery.go", "已实现"),
    ("身份文档与注册验签", "pkg/ath/identity.go；identity_test.go；handler/ath.go", "已实现并有负向测试"),
    ("双向握手与会话密钥", "pkg/ath/handshake.go；handshake_test.go", "已实现 ES256、ECDH、HKDF 与状态机"),
    ("请求完整性与重放保护", "pkg/ath/integrity.go；integrity_test.go", "已实现 HMAC、时间窗与 Redis nonce"),
    ("Token 与 Scope 绑定", "pkg/jwt；pkg/oauth；handler/ath.go", "已绑定 Agent、Provider、Session、Handshake"),
    ("签名审计链", "pkg/ath/audit.go；dao/ath_audit_records.go；init-db.sql", "已实现哈希链、签名和追加保护"),
    ("密钥轮换与远程签名", "pkg/ath/handshake.go；config；initApp.go", "已实现密钥环与 HTTPS KMS 网关"),
    ("Outbox 与外部锚定", "pkg/ath/anchor.go；dao/ath_audit_outbox.go", "已实现可靠投递、状态与重试"),
    ("自动化验证", "pkg/ath/*_test.go；scripts/test-ath.sh", "针对性测试通过，真实环境联调待完成"),
    ("MCP Server 与 CLI", "mcp/、cli/ 目录；README 路线图", "规划中"),
    ("多服务生态", "README 路线图中的文件、消息、权限服务", "规划中"),
], [1800, 5000, 2560])

add_heading(doc, "编制说明", 1)
add_para(doc, "本报告依据用户提供的《【报告大纲】智能体互联网治理_交互技术架.docx》编制，并结合 Agent Base 项目 README、交互治理实践材料及 user-service 代码进行事实核验。报告中的“已实现”“部分实现”“规划中”按当前仓库证据区分；行业应用与发展建议属于基于项目能力的研究判断，不代表相关能力已经完成生产部署。")

# Header/footer
for section in doc.sections:
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(hp.add_run("Agent Base | 智能体互联网交互治理研究报告"), size=8.5, color=GRAY)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(fp.add_run("Agent Base 开源项目团队 · 2026 年 6 月"), size=8.5, color=GRAY)

# Keep table rows from splitting where possible.
for table in doc.tables:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)

doc.core_properties.title = "智能体互联网治理交互技术架构与产业实践研究报告——基于 Agent Base 开源项目的治理实践"
doc.core_properties.subject = "Agent Base 智能体互联网交互治理"
doc.core_properties.author = "Agent Base 开源项目团队"
doc.core_properties.keywords = "智能体互联网, 交互治理, Agent Base, ATH, MCP, OAuth"
doc.save(OUT)
print(OUT.resolve())
